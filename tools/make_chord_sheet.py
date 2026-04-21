"""
Chord Sheet Generator

Generates a DOCX chord sheet from a JSON template.

Defaults
- Input: studio_master/song_template.json (repository root)
- Output: documents/<Title>_<Artist>_Key_<Key>.docx

Layout
- Header: "Title — Artist · BPM · Key" centered, with color accents for BPM/Key
- Sections: uppercase headings with icons and color accents
- Lines:
  - Dict line: chords line (Courier New, chords in bold), then lyrics line (Arial Black, 10pt)
  - String line: chords-only line (Courier New, chords in bold)
- Fade out support: render an italic "Fade out…" line when indicated
- Space saving: optional merge of adjacent short lyric lines

Usage (PowerShell)
  python .\$$make_chord_sheet.py               # uses defaults
  python .\$$make_chord_sheet.py --input studio_master\song_template.json --out documents

Note: Requires 'python-docx' (pip install python-docx)
"""

from __future__ import annotations

import argparse
import json
import os
import re
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple, Union

try:
	from docx import Document
	from docx.shared import Pt, RGBColor
	from docx.enum.text import WD_ALIGN_PARAGRAPH
except Exception as e:  # pragma: no cover
	raise SystemExit("python-docx is required. Install with: pip install python-docx") from e


JsonObj = Dict[str, Any]
Line = Union[str, Dict[str, Any]]


# Section style mapping: name contains key -> (icon, RGB color)
SECTION_STYLES: List[Tuple[re.Pattern, str, Tuple[int, int, int]]] = [
	(re.compile(r"^intro", re.I), "⭘", (90, 90, 90)),
	(re.compile(r"^verse", re.I), "♪", (30, 120, 200)),
	(re.compile(r"^pre[- ]?chorus", re.I), "↗", (120, 90, 200)),
	(re.compile(r"^chorus", re.I), "★", (200, 60, 60)),
	(re.compile(r"^bridge", re.I), "⇄", (130, 130, 130)),
	(re.compile(r"^interlude|solo", re.I), "◎", (120, 120, 120)),
	(re.compile(r"^outro", re.I), "◌", (90, 90, 90)),
]


# Heuristic to detect chord-like tokens (bold them)
CHORD_RE = re.compile(
	r"^(?:N|[A-G])(?:#|b)?(?:m|maj|min|dim|aug|sus[24]?|madd\d+|add\d+|M7|maj7|m7|6|5|7|9|11|13)?(?:/[A-G](?:#|b)?)?$"
)

# Visible separator used when merging adjacent short lyric lines
LYRICS_MERGE_SEPARATOR = " ‖ "  # U+2016 Double Vertical Line with spaces
# Recognize either unicode double bar or legacy ASCII "||" in lyrics when styling
_MERGE_SEP_PATTERN = re.compile(r"(\s*\|\|\s*|‖)")

# Visible separator between chord fragments when lyric lines are merged
CHORDS_MERGE_SEPARATOR = " || "

# Accent color for merge separators (lyrics + chords)
MERGE_SEP_COLOR: Tuple[int, int, int] = (200, 120, 40)


def is_chord_token(tok: str) -> bool:
	tok = tok.strip()
	if not tok:
		return False
	# Allow barlines/dots as separators but not as chords
	if tok in {"|", "||", "‖", "·", "•", "-", "—", "…", ".", "..", "..."}:
		return False
	return bool(CHORD_RE.match(tok))


def load_song(input_path: Optional[Path]) -> JsonObj:
	if input_path is None:
		input_path = Path(__file__).parent / "studio_master" / "song_template.json"
	if not input_path.exists():
		raise FileNotFoundError(f"Template not found: {input_path}")
	with input_path.open("r", encoding="utf-8") as f:
		data = json.load(f)
	# Basic normalization
	data.setdefault("title", "Untitled")
	data.setdefault("artist", "Unknown")
	data.setdefault("key", "?")
	data.setdefault("bpm", "?")
	data.setdefault("sections", [])
	return data


def sanitize_filename(name: str) -> str:
	keep = " _-.()[]{}+"
	return "".join(c for c in name if c.isalnum() or c in keep).strip()


def format_key_for_filename(key: str) -> str:
	"""Convert accidentals to file-safe text (e.g., F# -> F_Sharp, Bb -> B_Flat).

	Handles unicode accidentals (♯, ♭) and preserves the rest of the string.
	Examples:
	  - "F#" -> "F_Sharp"
	  - "Bb" -> "B_Flat"
	  - "C#m" -> "C_Sharp m"
	  - "Db Major" -> "D_Flat Major"
	"""
	if not isinstance(key, str):
		key = str(key)

	# Normalize unicode accidentals
	key_norm = key.replace("♯", "#").replace("♭", "b")

	def repl(match: re.Match) -> str:
		note = match.group(1)
		acc = match.group(2)
		return f"{note}_{'Sharp' if acc == '#' else 'Flat'}"

	# Replace first note+accidental occurrence; if multiple, replace all
	key_labeled = re.sub(r"([A-Ga-g])([#b])", repl, key_norm)
	return key_labeled


def compute_output_path(song: JsonObj, outdir: Optional[Path], lyrics_only: bool = False) -> Path:
	if outdir is None:
		outdir = Path(__file__).parent / "documents"
	outdir.mkdir(parents=True, exist_ok=True)
	title = sanitize_filename(song.get("title", "Untitled"))
	artist = sanitize_filename(song.get("artist", "Unknown"))
	raw_key = str(song.get("key", "?"))
	key_for_file = sanitize_filename(format_key_for_filename(raw_key))
	suffix = "_Lyrics_Only" if lyrics_only else ""
	filename = f"{title}_{artist}_Key_{key_for_file}{suffix}.docx"
	return outdir / filename


def add_header(doc: Document, song: JsonObj, *, small: bool = False) -> None:
	p = doc.add_paragraph()
	p.alignment = WD_ALIGN_PARAGRAPH.CENTER

	# Title — Artist · BPM · Key
	r_title = p.add_run(str(song.get("title", "Untitled")))
	r_title.bold = True
	r_title.font.size = Pt(12 if small else 16)

	p.add_run(" — ")

	r_artist = p.add_run(str(song.get("artist", "Unknown")))
	r_artist.font.size = Pt(10 if small else 14)

	p.add_run("  ·  ")

	r_bpm = p.add_run(f"BPM {song.get('bpm', '?')}")
	r_bpm.font.color.rgb = RGBColor(30, 120, 200)  # blue accent
	r_bpm.font.bold = True

	p.add_run("  ·  ")

	r_key = p.add_run(f"Key {song.get('key', '?')}")
	r_key.font.color.rgb = RGBColor(30, 160, 90)   # green accent
	r_key.font.bold = True

	# Top-level annotations: notes, vibe, arrangement_notes, etc.
	annotations_to_render = []
	if song.get("notes"):
		annotations_to_render.append(("📝", str(song.get("notes"))))
	if song.get("vibe"):
		annotations_to_render.append(("🎵", str(song.get("vibe"))))
	if song.get("arrangement_notes"):
		annotations_to_render.append(("🎼", str(song.get("arrangement_notes"))))
	
	for icon, text in annotations_to_render:
		p_ann = doc.add_paragraph()
		p_ann.alignment = WD_ALIGN_PARAGRAPH.CENTER
		r_ann = p_ann.add_run(f"{icon} {text}")
		r_ann.italic = True
		r_ann.font.size = Pt(7 if small else 9)
		r_ann.font.color.rgb = RGBColor(110, 110, 110)


def section_style(name: str) -> Tuple[str, Tuple[int, int, int]]:
	for pat, icon, rgb in SECTION_STYLES:
		if pat.search(name):
			return icon, rgb
	return "■", (100, 100, 100)


def add_section_heading(doc: Document, name: str, *, small: bool = False, annotation: Optional[str] = None) -> None:
	icon, rgb = section_style(name)
	p = doc.add_paragraph()
	p.paragraph_format.space_before = Pt(6 if small else 10)
	r = p.add_run(f"{icon} {name.upper()}")
	r.bold = True
	r.font.size = Pt(9 if small else 12)
	r.font.color.rgb = RGBColor(*rgb)
	
	# Add annotation inline to the right if provided
	if annotation:
		ann_lower = annotation.lower().strip()
		if ann_lower in ["walk up", "walkup"]:
			disp = "↗ Walkup"
		elif ann_lower in ["walk down", "walkdown"]:
			disp = "↓ Walkdown"
		elif ann_lower.startswith("rit") or "ritard" in ann_lower:
			disp = "Ritardando"
		else:
			disp = annotation
		
		r_ann = p.add_run(f"  {disp}")
		r_ann.italic = True
		r_ann.font.size = Pt(7 if small else 9)
		r_ann.font.color.rgb = RGBColor(120, 80, 160)


def add_chords_paragraph(doc: Document, text: str, *, small: bool = False) -> None:
	p = doc.add_paragraph()
	# Split while keeping whitespace to preserve spacing, and also detect '||' groups
	for tok in re.split(r"(\s+)", text):
		if not tok:
			continue
		if tok.strip() == "||":
			r = p.add_run(CHORDS_MERGE_SEPARATOR)
			r.font.name = "Courier New"
			r.font.size = Pt(8 if small else 10)
			r.bold = True
			r.font.color.rgb = RGBColor(*MERGE_SEP_COLOR)
			continue
		r = p.add_run(tok)
		r.font.name = "Courier New"
		r.font.size = Pt(8 if small else 10)
		if is_chord_token(tok):
			r.bold = True


def add_lyrics_paragraph(doc: Document, text: str, *, small: bool = False, backing_vocals: Optional[str] = None) -> None:
	"""Render lyrics and style merge separators so breaks are visually clear.
	
	If backing_vocals is provided and the combined length is reasonable, append inline with separator.
	Otherwise render backing vocals on a separate line beneath.
	"""
	p = doc.add_paragraph()
	parts = _MERGE_SEP_PATTERN.split(text)
	for part in parts:
		if part is None or part == "":
			continue
		if _MERGE_SEP_PATTERN.fullmatch(part):
			# Normalize any form to our preferred separator and style subtly
			r = p.add_run(LYRICS_MERGE_SEPARATOR)
			r.italic = True
			r.font.name = "Arial Black"
			r.font.size = Pt(8 if small else 10)
			r.font.color.rgb = RGBColor(*MERGE_SEP_COLOR)
			continue
		r = p.add_run(part)
		r.font.name = "Arial Black"
		r.font.size = Pt(8 if small else 10)
	
	# Add backing vocals inline if reasonable length, otherwise on separate line
	if backing_vocals and backing_vocals.strip():
		combined_len = len(text) + len(backing_vocals)
		# If combined length is reasonable (< 100 chars), render inline
		if combined_len < 100:
			# Add separator
			r_sep = p.add_run("  •  ")
			r_sep.font.name = "Arial Black"
			r_sep.font.size = Pt(7 if small else 9)
			r_sep.font.color.rgb = RGBColor(90, 110, 150)
			# Add backing vocals inline
			r_bv = p.add_run(backing_vocals.strip())
			r_bv.italic = True
			r_bv.font.name = "Arial"
			r_bv.font.size = Pt(7 if small else 9)
			r_bv.font.color.rgb = RGBColor(90, 110, 150)
		else:
			# Render on separate line with minimal spacing
			p.paragraph_format.space_after = Pt(0)
			p_bv = doc.add_paragraph()
			p_bv.paragraph_format.space_before = Pt(0)
			p_bv.paragraph_format.space_after = Pt(2 if small else 3)
			r_bv = p_bv.add_run(f"[BVs] {backing_vocals.strip()}")
			r_bv.italic = True
			r_bv.font.name = "Arial"
			r_bv.font.size = Pt(7 if small else 9)
			r_bv.font.color.rgb = RGBColor(90, 110, 150)


def add_annotation_paragraph(doc: Document, text: str, *, kind: str = "conductor", small: bool = False) -> None:
	"""Render a small italic annotation line for conductor cues (e.g., Ritardando, Walkup) or notes."""
	color_map = {
		"conductor": (120, 80, 160),
		"note": (110, 110, 110),
		"backing": (90, 110, 150),
	}
	rgb = color_map.get(kind, color_map["conductor"])
	p = doc.add_paragraph()
	r = p.add_run(text)
	r.italic = True
	r.font.size = Pt(7 if small else 9)
	r.font.color.rgb = RGBColor(*rgb)


def add_backing_vocals_paragraph(doc: Document, text: str, *, small: bool = False) -> None:
	"""Render backing vocals as a subtle, annotated line beneath the main lyrics."""
	add_annotation_paragraph(doc, f"[BackingVox] {text}", kind="backing", small=small)


# ---------- Repeat-link detection utilities ----------
def _base_section_name(name: str) -> str:
	"""Normalize section name to a base label (e.g., 'Chorus 2' -> 'chorus')."""
	n = name.strip().lower()
	# remove trailing numerals and parentheses like "(2)"
	n = re.sub(r"\(\s*\d+\s*\)$", "", n)
	n = re.sub(r"\s*\d+\s*$", "", n)
	# collapse whitespace
	n = re.sub(r"\s+", " ", n)
	return n.strip()


def _normalize_text(s: str, keep_letters_digits_only: bool = True) -> str:
	s = s.lower().strip()
	if keep_letters_digits_only:
		s = re.sub(r"[^a-z0-9\s]", "", s)
	s = re.sub(r"\s+", " ", s)
	return s


def _collect_section_texts(section: JsonObj) -> Tuple[str, str]:
	"""Return concatenated (chords_text, lyrics_text) for a section."""
	chords_parts: List[str] = []
	lyrics_parts: List[str] = []
	for line in section.get("lines", []) or []:
		if isinstance(line, str):
			chords_parts.append(line)
			continue
		if isinstance(line, dict):
			if line.get("chords"):
				chords_parts.append(str(line.get("chords")))
			if line.get("lyrics"):
				lyrics_parts.append(str(line.get("lyrics")))
	chords_text = _normalize_text(" ".join(chords_parts), keep_letters_digits_only=False)
	lyrics_text = _normalize_text(" ".join(lyrics_parts), keep_letters_digits_only=True)
	return chords_text, lyrics_text


def _similarity_ratio(a: str, b: str) -> float:
	"""Compute a quick similarity ratio using SequenceMatcher-like approach without import."""
	# Lightweight token Jaccard as a dependency-free metric
	ta = set(a.split())
	tb = set(b.split())
	if not ta and not tb:
		return 1.0
	if not ta or not tb:
		return 0.0
	inter = len(ta & tb)
	union = len(ta | tb)
	return inter / union if union else 0.0


def _collect_raw_lyrics(section: JsonObj) -> str:
	"""Collect raw (un-normalized) lyrics from a section, concatenated compactly."""
	parts: List[str] = []
	for line in section.get("lines", []) or []:
		if isinstance(line, dict):
			ly = str(line.get("lyrics", "")).strip()
			if ly:
				parts.append(ly)
	# Use a compact separator to save space
	return "  ·  ".join(parts)


def _chord_tokens(text: str) -> List[str]:
	"""Extract only chord-like tokens from a normalized text, preserving order."""
	return [t for t in text.split() if is_chord_token(t)]


def _merge_chords(a: str, b: str) -> str:
	"""Join two chord strings with a double bar, normalizing stray bars/spaces.

	Removes a trailing single bar from `a` and a leading single bar from `b`
	to avoid artifacts like "| || |".
	"""
	a = str(a or "").strip()
	b = str(b or "").strip()
	if not a:
		return b
	if not b:
		return a
	a = re.sub(r"\s*\|\s*$", "", a)
	b = re.sub(r"^\s*\|\s*", "", b)
	return f"{a}{CHORDS_MERGE_SEPARATOR}{b}"


def _set_jaccard(a: List[str] | set, b: List[str] | set) -> float:
	sa = set(a)
	sb = set(b)
	if not sa and not sb:
		return 1.0
	if not sa or not sb:
		return 0.0
	return len(sa & sb) / max(1, len(sa | sb))


def _bigram_jaccard(a_tokens: List[str], b_tokens: List[str]) -> float:
	def bigrams(seq: List[str]) -> set:
		return set(zip(seq, seq[1:])) if len(seq) >= 2 else set()
	return _set_jaccard(bigrams(a_tokens), bigrams(b_tokens))


def _sections_are_similar(
	a: JsonObj,
	b: JsonObj,
	*,
	chord_threshold: float,
	lyric_threshold: float,
	use_lyrics: bool,
) -> Tuple[bool, float, float, float]:
	"""Compare sections with chord-first logic.

	Returns (ok, chord_sim, lyric_sim, chord_bigram_sim).
	"""
	a_ch, a_ly = _collect_section_texts(a)
	b_ch, b_ly = _collect_section_texts(b)
	# Chord similarity: combine unigram and bigram Jaccard over chord tokens
	a_ct = _chord_tokens(a_ch)
	b_ct = _chord_tokens(b_ch)
	uni = _set_jaccard(a_ct, b_ct)
	bi = _bigram_jaccard(a_ct, b_ct)
	# If sequences are short, lean more on unigrams
	ch_sim = uni if len(a_ct) < 2 or len(b_ct) < 2 else (0.6 * bi + 0.4 * uni)
	ly_sim = _similarity_ratio(a_ly, b_ly)
	if use_lyrics:
		ok = (ch_sim >= chord_threshold) and (ly_sim >= lyric_threshold)
	else:
		ok = (ch_sim >= chord_threshold)
	return ok, ch_sim, ly_sim, bi


def link_repeated_sections(
	sections: List[JsonObj],
	threshold: float = 0.95,
	same_name_only: bool = True,
	*,
	chord_threshold: Optional[float] = None,
	lyric_threshold: float = 0.2,
	use_lyrics: bool = False,
	only_chorus: bool = False,
) -> List[JsonObj]:
	"""Mark sections that are near-duplicates of previous ones.

	Returns a new list of sections. Any linked section will have a special
	key "__link_to" with the name of the earlier section to reference.
	"""
	out: List[JsonObj] = []
	history: List[Tuple[str, JsonObj]] = []  # (base_name, section_obj)
	ch_th = threshold if chord_threshold is None else float(chord_threshold)
	for sec in sections or []:
		name = str(sec.get("name", "")).strip() or "SECTION"
		base = _base_section_name(name)

		# Skip linking if only_chorus is True and this isn't a chorus
		if only_chorus and "chorus" not in base.lower():
			out.append(sec)
			continue

		link_name: Optional[str] = None
		# Try to find a previous section to link to
		for prev_base, prev_sec in history:
			if same_name_only and prev_base != base:
				continue
			ok, ch_sim, ly_sim, _ = _sections_are_similar(
				sec,
				prev_sec,
				chord_threshold=ch_th,
				lyric_threshold=float(lyric_threshold),
				use_lyrics=bool(use_lyrics),
			)
			if ok:
				# use the previous section's exact name for user-facing reference
				link_name = str(prev_sec.get("name", name))
				break

		if link_name:
			shallow = dict(sec)
			shallow["__link_to"] = link_name
			out.append(shallow)
		else:
			out.append(sec)
			history.append((base, sec))
	return out


def merge_short_lyric_lines(lines: List[Line], threshold: int = 70) -> List[Line]:
	"""Merge adjacent short lyric dict lines and insert a visible separator.

	- Keeps merging while combined length (including separator) stays under threshold.
	- Preserves and concatenates chords as previously (joined with " | ").
	- Inserts LYRICS_MERGE_SEPARATOR between lyric fragments.
	- Preserves backing_vocals, notes, annotation, and repeat from first line only.
	"""
	out: List[Line] = []
	i = 0
	while i < len(lines):
		cur = lines[i]
		if (
			isinstance(cur, dict)
			and "lyrics" in cur
			and isinstance(cur["lyrics"], str)
		):
			lyric = cur.get("lyrics", "")
			chords = cur.get("chords", "")
			# Preserve other fields from the first line
			backing_vocals = cur.get("backing_vocals")
			notes = cur.get("notes")
			annotation = cur.get("annotation")
			repeat = cur.get("repeat")
			j = i + 1
			merged = False
			while j < len(lines):
				nxt = lines[j]
				if not (isinstance(nxt, dict) and isinstance(nxt.get("lyrics", None), str)):
					break
				next_ly = nxt.get("lyrics", "").strip()
				# Compute prospective length including separator
				sep = "" if not lyric else LYRICS_MERGE_SEPARATOR
				if len(lyric) + len(sep) + len(next_ly) <= threshold:
					# Avoid doubling separators if user already typed one at the end
					if re.search(r"(\|\|\s*$|‖\s*$)", lyric):
						lyric = lyric.rstrip() + " " + next_ly
					else:
						lyric = (lyric + sep + next_ly) if lyric else next_ly
					if nxt.get("chords"):
						chords = _merge_chords(chords, str(nxt["chords"]))
					j += 1
					merged = True
				else:
					break
			# Build result dict with all preserved fields
			result = {"lyrics": lyric, "chords": chords}
			if backing_vocals is not None:
				result["backing_vocals"] = backing_vocals
			if notes is not None:
				result["notes"] = notes
			if annotation is not None:
				result["annotation"] = annotation
			if repeat is not None:
				result["repeat"] = repeat
			out.append(result)
			i = j if merged else i + 1
		else:
			out.append(cur)
			i += 1
	return out


def render_section(doc: Document, section: JsonObj, do_merge: bool = True, *, hide_chords: bool = False, small: bool = False) -> None:
	name = str(section.get("name", "")).strip() or "SECTION"
	
	# Check for section-level repeat to include in heading
	try:
		sec_repeat = int(section.get("repeat") or section.get("repeat_times") or 0)
	except Exception:
		sec_repeat = 0
	
	# Add section heading with optional repeat notation
	if sec_repeat and sec_repeat > 1:
		icon, rgb = section_style(name)
		p = doc.add_paragraph()
		p.paragraph_format.space_before = Pt(6 if small else 10)
		r = p.add_run(f"{icon} {name.upper()} ")
		r.bold = True
		r.font.size = Pt(9 if small else 12)
		r.font.color.rgb = RGBColor(*rgb)
		r_repeat = p.add_run(f"Repeat ×{sec_repeat}")
		r_repeat.bold = True
		r_repeat.font.size = Pt(9 if small else 12)
		r_repeat.font.color.rgb = RGBColor(200, 120, 40)  # Orange accent for repeat
		# Add annotation inline if present
		sec_ann_str = section.get("annotation") or section.get("annotations") or section.get("notes")
		if isinstance(sec_ann_str, str) and sec_ann_str.strip():
			ann_lower = sec_ann_str.lower().strip()
			if ann_lower in ["walk up", "walkup"]:
				disp = "↗ Walkup"
			elif ann_lower in ["walk down", "walkdown"]:
				disp = "↓ Walkdown"
			elif ann_lower.startswith("rit") or "ritard" in ann_lower:
				disp = "Ritardando"
			else:
				disp = sec_ann_str.strip()
			r_ann = p.add_run(f"  {disp}")
			r_ann.italic = True
			r_ann.font.size = Pt(7 if small else 9)
			r_ann.font.color.rgb = RGBColor(120, 80, 160)
	else:
		# Check for annotation to pass to heading
		sec_ann_str = section.get("annotation") or section.get("annotations") or section.get("notes")
		if isinstance(sec_ann_str, str):
			add_section_heading(doc, name, small=small, annotation=sec_ann_str.strip())
		else:
			add_section_heading(doc, name, small=small)

	# Section-level annotations/notes/backing vocals (render under heading) - only for lists
	sec_ann = section.get("annotation") or section.get("annotations") or section.get("notes")
	if isinstance(sec_ann, list):
		# Render multiple annotations if provided as a list
		for ann in sec_ann:
			try:
				ann_text = str(ann).strip()
			except Exception:
				ann_text = ""
			if ann_text:
				ann_lower = ann_text.lower()
				if ann_lower in ["walk up", "walkup"]:
					add_annotation_paragraph(doc, "↗ Walkup", kind="conductor", small=small)
				elif ann_lower in ["walk down", "walkdown"]:
					add_annotation_paragraph(doc, "↓ Walkdown", kind="conductor", small=small)
				elif ann_lower.startswith("rit") or "ritard" in ann_lower:
					add_annotation_paragraph(doc, "Ritardando", kind="conductor", small=small)
				elif "harmony" in ann_lower or "harmonies" in ann_lower:
					add_annotation_paragraph(doc, ann_text, kind="backing", small=small)
				else:
					add_annotation_paragraph(doc, ann_text, kind="note", small=small)

	sec_bv = section.get("backing_vocals")
	if sec_bv:
		if isinstance(sec_bv, list):
			bv_text = "  ·  ".join(str(x) for x in sec_bv if str(x).strip())
		else:
			bv_text = str(sec_bv)
		if bv_text.strip():
			add_backing_vocals_paragraph(doc, bv_text.strip(), small=small)

	# Get raw lines before checking for link (needed for compact chord display)
	raw_lines: List[Line] = section.get("lines", []) or []

	# If this section is a link to an earlier section, render a short reference and return
	link_to = section.get("__link_to")
	if link_to:
		p = doc.add_paragraph()
		r = p.add_run(f"Same as {str(link_to).upper()} (see above)")
		r.italic = True
		r.font.size = Pt(8 if small else 10)
		
		# Extract and render compact chord progression
		compact_chords = []
		for line in raw_lines:
			if isinstance(line, dict) and line.get("chords"):
				ch = str(line.get("chords", "")).strip()
				if ch:
					compact_chords.append(ch)
		
		if compact_chords:
			# Render compact progression on same line or next line depending on length
			prog_text = "  ·  ".join(compact_chords)
			if len(prog_text) < 80:
				# Append to same line
				r_sep = p.add_run("  —  ")
				r_sep.font.size = Pt(7 if small else 9)
				r_sep.font.color.rgb = RGBColor(110, 110, 110)
				r_prog = p.add_run(prog_text)
				r_prog.font.name = "Courier New"
				r_prog.font.size = Pt(7 if small else 9)
				r_prog.font.color.rgb = RGBColor(110, 110, 110)
			else:
				# Render on separate line
				p_prog = doc.add_paragraph()
				p_prog.paragraph_format.space_before = Pt(0)
				r_prog = p_prog.add_run(prog_text)
				r_prog.font.name = "Courier New"
				r_prog.font.size = Pt(7 if small else 9)
				r_prog.italic = True
				r_prog.font.color.rgb = RGBColor(110, 110, 110)
		
		# Also show fine-print lyrics for convenience
		fp = _collect_raw_lyrics(section).strip()
		if fp:
			p_fp = doc.add_paragraph()
			p_fp.paragraph_format.space_before = Pt(1 if small else 2)
			r_fp = p_fp.add_run(fp)
			r_fp.italic = True
			r_fp.font.size = Pt(6 if small else 7)
			r_fp.font.color.rgb = RGBColor(110, 110, 110)
		# If this linked section is to be repeated, still show the repeat note
		try:
			sec_repeat = int(section.get("repeat") or section.get("repeat_times") or 0)
		except Exception:
			sec_repeat = 0
		if sec_repeat and sec_repeat > 1:
			p2 = doc.add_paragraph()
			r2 = p2.add_run(f"Repeat {name.upper()} ×{sec_repeat}")
			r2.bold = True
			r2.font.size = Pt(8 if small else 10)
		return

	lines = merge_short_lyric_lines(raw_lines) if do_merge else raw_lines

	for line in lines:
		# Fade-out sentinel on a line-as-dict
		if isinstance(line, dict) and line.get("fade_out"):
			# Show a subtle note even if chords are hidden
			add_annotation_paragraph(doc, "Fade out…", kind="note", small=small)
			continue

		if isinstance(line, str):
			# chords-only
			if not hide_chords:
				add_chords_paragraph(doc, line, small=small)
			continue

		if isinstance(line, dict):
			# dict line: chords then lyrics
			chords = str(line.get("chords", "")).strip()
			lyrics = str(line.get("lyrics", "")).strip()
			notes = str(line.get("notes", "")).strip()
			annotation = str(line.get("annotation", "")).strip()
			repeat = line.get("repeat")
			backing = line.get("backing_vocals")
			
			# Process backing vocals (can be string or list)
			backing_text = None
			if backing:
				if isinstance(backing, list):
					backing_text = "  ·  ".join(str(x) for x in backing if str(x).strip())
				else:
					backing_text = str(backing)
				if not backing_text.strip():
					backing_text = None

			# Optional conductor/notation annotation line
			if annotation:
				ann = annotation.strip().strip("()")
				ann_lower = ann.lower()
				if ann_lower in ["walk up", "walkup"]:
					ann_disp = "↗ Walkup"
				elif ann_lower in ["walk down", "walkdown"]:
					ann_disp = "↓ Walkdown"
				elif ann_lower.startswith("rit") or "ritard" in ann_lower:
					ann_disp = "Ritardando"
				else:
					ann_disp = ann
				add_annotation_paragraph(doc, ann_disp, kind="conductor", small=small)

			if chords and not hide_chords:
				add_chords_paragraph(doc, chords, small=small)
			if lyrics:
				add_lyrics_paragraph(doc, lyrics, small=small, backing_vocals=backing_text)
			# Notes remain separate (structural annotations, not harmony)
			if notes:
				add_annotation_paragraph(doc, notes, kind="note", small=small)
			if repeat:
				p = doc.add_paragraph()
				r = p.add_run(f"×{repeat}")
				r.bold = True
				r.font.size = Pt(7 if small else 9)
			continue

		# unknown type -> string cast (skip if hiding chords)
		if not hide_chords:
			add_chords_paragraph(doc, str(line), small=small)


def build_docx(song: JsonObj, outpath: Path, do_merge: bool = True, *, hide_chords: bool = False, small: bool = False) -> Path:
	doc = Document()

	add_header(doc, song, small=small)

	sections = song.get("sections", []) or []
	# Optionally link repeated sections if enabled in the song dict
	link_cfg = song.get("link_repeats")  # can be bool or dict
	link_enabled = False
	link_threshold = 0.95
	link_same_name_only = True
	link_chord_threshold: Optional[float] = None
	link_lyric_threshold = 0.2
	link_use_lyrics = False
	link_only_chorus = False
	if isinstance(link_cfg, bool):
		link_enabled = link_cfg
	elif isinstance(link_cfg, dict):
		link_enabled = bool(link_cfg.get("enabled", True))
		link_threshold = float(link_cfg.get("threshold", link_threshold))
		link_same_name_only = bool(link_cfg.get("same_name_only", link_same_name_only))
		# Optional fine-tuning knobs
		if "chord_threshold" in link_cfg:
			try:
				link_chord_threshold = float(link_cfg.get("chord_threshold"))
			except Exception:
				link_chord_threshold = None
		if "lyric_threshold" in link_cfg:
			try:
				link_lyric_threshold = float(link_cfg.get("lyric_threshold"))
			except Exception:
				pass
		if "ignore_lyrics" in link_cfg:
			link_use_lyrics = not bool(link_cfg.get("ignore_lyrics"))
		if "only_chorus" in link_cfg:
			link_only_chorus = bool(link_cfg.get("only_chorus"))

	if getattr(build_docx, "_cli_link_repeats", None) is not None:
		# CLI flag takes precedence over JSON
		link_enabled = bool(getattr(build_docx, "_cli_link_repeats"))
		link_threshold = float(getattr(build_docx, "_cli_link_threshold", link_threshold))
		link_same_name_only = bool(getattr(build_docx, "_cli_link_same_name_only", link_same_name_only))
		# CLI overrides for new knobs
		cli_ch = getattr(build_docx, "_cli_link_chord_threshold", None)
		if cli_ch is not None:
			link_chord_threshold = float(cli_ch)
		cli_ly = getattr(build_docx, "_cli_link_lyric_threshold", None)
		if cli_ly is not None:
			link_lyric_threshold = float(cli_ly)
		cli_ignore = getattr(build_docx, "_cli_link_ignore_lyrics", None)
		if cli_ignore is not None:
			link_use_lyrics = not bool(cli_ignore)
		cli_only_chorus = getattr(build_docx, "_cli_link_only_chorus", None)
		if cli_only_chorus is not None:
			link_only_chorus = bool(cli_only_chorus)

	if link_enabled:
		sections = link_repeated_sections(
			sections,
			threshold=link_threshold,
			same_name_only=link_same_name_only,
			chord_threshold=link_chord_threshold,
			lyric_threshold=link_lyric_threshold,
			use_lyrics=link_use_lyrics,
			only_chorus=link_only_chorus,
		)
	for sec in sections:
		render_section(doc, sec, do_merge=do_merge, hide_chords=hide_chords, small=small)

	# Document metadata
	core = doc.core_properties
	core.title = str(song.get("title", ""))
	core.author = str(song.get("artist", ""))
	core.comments = "Generated by chord sheet generator"

	doc.save(str(outpath))
	return outpath


def parse_args() -> argparse.Namespace:
	p = argparse.ArgumentParser(description="Generate a DOCX chord sheet from a JSON template")
	p.add_argument("--input", "-i", type=str, default=None, help="Path to song JSON (default: studio_master/song_template.json)")
	p.add_argument("--out", "-o", type=str, default=None, help="Output directory (default: documents)")
	p.add_argument("--no-merge", action="store_true", help="Disable merging of short lyric lines")
	p.add_argument("--link-repeats", action="store_true", default=True, help="Symbolically link sections that repeat (dedupe repeated sections)")
	p.add_argument("--link-threshold", type=float, default=0.95, help="Legacy combined threshold (used as chord threshold unless overridden)")
	p.add_argument("--link-chord-threshold", type=float, default=None, help="Chord similarity threshold (overrides --link-threshold)")
	p.add_argument("--link-lyrics-threshold", type=float, default=0.2, help="Lyrics similarity threshold (default low)")
	p.add_argument("--link-ignore-lyrics", action="store_true", help="Ignore lyrics similarity when linking (chord-first)")
	p.add_argument("--link-any-name", action="store_true", help="Allow linking sections across different names (e.g., Verse to Chorus)")
	p.add_argument("--link-only-chorus", action="store_true", help="Only link chorus sections (for lyrics-only sheets)")
	# Output controls
	p.add_argument("--no-chords", action="store_true", help="Omit chord lines (lyrics-only output)")
	p.add_argument("--lyrics-only", action="store_true", help="Alias for --no-chords")
	p.add_argument("--small", action="store_true", help="Compact mode with smaller fonts (for instrumentalists)")
	return p.parse_args()


def main() -> None:
	args = parse_args()
	input_path = Path(args.input) if args.input else None
	outdir = Path(args.out) if args.out else None

	song = load_song(input_path)
	hide = bool(getattr(args, "no_chords", False) or getattr(args, "lyrics_only", False))
	outpath = compute_output_path(song, outdir, lyrics_only=hide)
	# Expose CLI preferences to build_docx via attributes (keeps function signature stable)
	setattr(build_docx, "_cli_link_repeats", args.link_repeats)
	setattr(build_docx, "_cli_link_threshold", args.link_threshold)
	setattr(build_docx, "_cli_link_same_name_only", not args.link_any_name)
	setattr(build_docx, "_cli_link_chord_threshold", args.link_chord_threshold)
	setattr(build_docx, "_cli_link_lyric_threshold", args.link_lyrics_threshold)
	setattr(build_docx, "_cli_link_ignore_lyrics", args.link_ignore_lyrics)
	setattr(build_docx, "_cli_link_only_chorus", args.link_only_chorus)
	small = bool(getattr(args, "small", False))
	build_docx(song, outpath, do_merge=not args.no_merge, hide_chords=hide, small=small)
	print(f"Saved: {outpath}")


if __name__ == "__main__":
	main()

